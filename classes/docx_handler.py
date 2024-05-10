import tempfile
import subprocess
import docx
import re
import uuid
import datetime
import os

from .document_handler import DocumentHandler
from python_redlines.engines import XmlPowerToolsEngine


class DOCXHandler(DocumentHandler):

    def __init__(self, file_path):
        super().__init__(file_path)
        self.file_path = file_path
        self.temp_file = None
        self.document = docx.Document(file_path)
        self.wrapper = XmlPowerToolsEngine()

    def count_pages(self):
        # This is a simplistic way of counting pages that may not be accurate for all documents.
        return sum(1 for p in self.document.paragraphs if p.text == "\f") + 1

    def remove_unwanted_pages(self, doc, keep_start, keep_end):
        """
        Removes pages not within the specified range.
        Pages are identified by manual page breaks or section breaks.
        """
        current_page = 1
        elements_to_keep = []
        for element in doc.element.body:
            if current_page >= keep_start and (keep_end is None or current_page <= keep_end):
                elements_to_keep.append(element)
            if element.tag.endswith('br') and element.get('type') == 'page':
                current_page += 1
                if current_page > keep_end:
                    break

        # Clear the document
        while doc.element.body.hasChildNodes():
            doc.element.body.removeChild(doc.element.body.lastChild)

        # Re-add the elements to keep
        for element in elements_to_keep:
            doc.element.body.appendChild(element)
    
    def get_first_n_pages(self, n):
        return self.transcribe_page_content(1, n)

    def get_last_n_pages(self, n):
        total_pages = self.count_pages()
        start_page = max(total_pages - n + 1, 1)
        return self.transcribe_page_content(start_page, total_pages)

    def get_page_n(self, n):
        return self.transcribe_page_content(n, n)
    
    def transcribe(self):

        # Convert the temporary DOCX file to Markdown using Pandoc
        with tempfile.NamedTemporaryFile(delete=False, suffix=".md") as temp_md:
            pandoc_command = [
                "pandoc",
                "-f", "docx",
                "-t", "markdown_strict",
                self.file_path,
                "-o", temp_md.name
            ]
            subprocess.run(pandoc_command, check=True)

        return temp_md.name
    
    def add_comment(self, input_string, comment_text, author='Essence of JJ', initials='AI'):

        # Loop through paragraphs and runs to find the input string
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                if input_string in run.text:
                    run.add_comment(comment_text, author=author, initials=initials)
                    self.document.save(self.file_path)
                    return True
                # Check paragraph as last recsort
                if input_string in paragraph.text:
                    paragraph.add_comment(comment_text, author=author, initials=initials)
                    self.document.save(self.file_path)
                    return True
        return False
    
    def track_change(self, input_string, replacement_string, author="Essence of JJ"):

        # Flag to track if a replacement was made
        replacement_made = False
            
        # Create a temporary file to save the modified document
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            # Create a copy of the original document in the temporary file
            temp_doc = docx.Document(self.file_path)

            # Loop through paragraphs
            for paragraph in temp_doc.paragraphs:
                # Join all runs in the paragraph into a single string
                paragraph_text = ''.join(run.text for run in paragraph.runs)

                # Use regular expressions to replace the input string
                new_paragraph_text = re.sub(re.escape(input_string), replacement_string, paragraph_text)

                # Check if a replacement was made
                if new_paragraph_text != paragraph_text:
                    replacement_made = True

                    # Clear the paragraph and add the new text
                    paragraph.clear()
                    paragraph.add_run(new_paragraph_text)

            if replacement_made:
                # Save the modified document to the temporary file
                temp_doc.save(temp_file.name)

        if replacement_made:
            # Run redline comparison
            output = self.wrapper.run_redline(author, self.file_path, temp_file.name)
            with open(self.file_path, 'wb') as f:
                f.write(output[0])

