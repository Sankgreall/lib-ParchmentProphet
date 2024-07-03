import tempfile
import subprocess
import docx
import re
import uuid
import datetime
import os
from docx.shared import RGBColor


from ..document_handler import DocumentHandler
from python_redlines.engines import XmlPowerToolsEngine


class DOCXHandler(DocumentHandler):

    def __init__(self, file_path):
        super().__init__(file_path)
        self.file_path = file_path
        self.temp_file = None
        self.document = docx.Document(file_path)
        self.wrapper = XmlPowerToolsEngine()
    
    def transcribe(self):

        # Convert the temporary DOCX file to Markdown using Pandoc
        with tempfile.NamedTemporaryFile(delete=False, suffix=".md") as temp_md:
            pandoc_command = [
                "pandoc",
                "-f", "docx",
                "-t", "markdown_strict",
                self.file_path,
                "-o", temp_md.name
            ]
            subprocess.run(pandoc_command, check=True)

        return temp_md.name
    
    def add_comment(self, input_string, comment_text, author='Essence of JJ', initials='AI'):

        # Loop through paragraphs and runs to find the input string
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                if input_string in run.text:
                    run.add_comment(comment_text, author=author, initials=initials)
                    self.document.save(self.file_path)
                    return True
                # Check paragraph as last recsort
                if input_string in paragraph.text:
                    paragraph.add_comment(comment_text, author=author, initials=initials)
                    self.document.save(self.file_path)
                    return True
        return False
    
    def track_change(self, input_string, replacement_string, author="Essence of JJ"):

        # Flag to track if a replacement was made
        replacement_made = False
            
        # Create a temporary file to save the modified document
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            # Create a copy of the original document in the temporary file
            temp_doc = docx.Document(self.file_path)

            # Loop through paragraphs
            for paragraph in temp_doc.paragraphs:
                # Join all runs in the paragraph into a single string
                paragraph_text = ''.join(run.text for run in paragraph.runs)

                # Use regular expressions to replace the input string
                new_paragraph_text = re.sub(re.escape(input_string), replacement_string, paragraph_text)

                # Check if a replacement was made
                if new_paragraph_text != paragraph_text:
                    replacement_made = True

                    # Clear the paragraph and add the new text
                    paragraph.clear()
                    paragraph.add_run(new_paragraph_text)

            if replacement_made:
                # Save the modified document to the temporary file
                temp_doc.save(temp_file.name)

        if replacement_made:
            # Run redline comparison
            output = self.wrapper.run_redline(author, self.file_path, temp_file.name)
            with open(self.file_path, 'wb') as f:
                f.write(output[0])
        

    def convert_markdown_to_paragraphs(self, markdown_text, custom_style=None, style_mapping=None):
        if style_mapping is None:
            style_mapping = {
                'heading1': 'Heading 1',
                'heading2': 'Heading 2',
                'bullet': 'List Bullet',
                'body': 'Normal'
            }

        paragraphs = []
        lines = markdown_text.split('\n')
        for line in lines:
            if line.startswith('# '):
                # Heading 1
                text = line[2:]
                style = custom_style or style_mapping.get('heading1', 'Heading 1')
                paragraphs.append((text, style))
            elif line.startswith('## '):
                # Heading 2
                text = line[3:]
                style = custom_style or style_mapping.get('heading2', 'Heading 2')
                paragraphs.append((text, style))
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                text = line[2:]
                style = custom_style or style_mapping.get('bullet', 'List Bullet')  # Adjust the style as needed
                paragraphs.append((text, style))
            else:
                # Paragraph
                style = custom_style or style_mapping.get('body', 'Normal')
                # Replace markdown syntax with custom tags
                text = re.sub(r'__(.*?)__', r'<underline>\1</underline>', line)  # Underline
                text = re.sub(r'\*\*(.*?)\*\*', r'<bold>\1</bold>', text)  # Bold
                text = re.sub(r'_(.*?)_', r'<italic>\1</italic>', text)   # Italics
                paragraphs.append((text, style))

        return paragraphs

    def apply_styles(self, paragraph, paragraphs, custom_color=None):
        # Clear existing content
        for run in paragraph.runs:
            run.clear()

        for i, (text, style_name) in enumerate(paragraphs):
            if i == 0:
                # First paragraph, modify in-place
                self.add_text_with_formatting(paragraph, text, style_name, custom_color)
            else:
                # Additional paragraphs
                new_paragraph = paragraph.insert_paragraph_before('')
                new_paragraph.style = self.document.styles[style_name]
                self.add_text_with_formatting(new_paragraph, text, style_name, custom_color)

    def add_text_with_formatting(self, paragraph, text, style_name, custom_color=None):
        paragraph.style = self.document.styles[style_name]
        parts = re.split(r'(<bold>|</bold>|<italic>|</italic>|<underline>|</underline>)', text)
        is_bold = False
        is_italic = False
        is_underline = False
        for part in parts:
            if part == '<bold>':
                is_bold = True
            elif part == '</bold>':
                is_bold = False
            elif part == '<italic>':
                is_italic = True
            elif part == '</italic>':
                is_italic = False
            elif part == '<underline>':
                is_underline = True
            elif part == '</underline>':
                is_underline = False
            else:
                run = paragraph.add_run(part)
                run.bold = is_bold
                run.italic = is_italic
                run.underline = is_underline
                if custom_color:
                    run.font.color.rgb = RGBColor(*custom_color)

    def insert_text_at_placeholder(self, placeholder, markdown_text, custom_style=None, custom_color=None, style_mapping=None):
        # Convert Markdown to paragraphs
        paragraphs = self.convert_markdown_to_paragraphs(markdown_text, custom_style, style_mapping)
        print("Paragraphs:", paragraphs)

        # Concatenate the formatted text into a single string
        replacement_text = "\n".join([text for text, _ in paragraphs])

        # Loop through paragraphs in the document
        for paragraph in self.document.paragraphs:
            if placeholder in paragraph.text:
                # Perform the replacement
                paragraph.text = paragraph.text.replace(placeholder, replacement_text)
                
                # Apply styles to paragraphs
                self.apply_styles(paragraph, paragraphs, custom_color)

                self.document.save(self.file_path)
                return True

        # If the placeholder was not found in paragraphs, check in tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        # Perform the replacement
                        cell_text = cell.text.replace(placeholder, replacement_text)
                        cell.paragraphs[0].text = cell_text
                        
                        # Apply styles to paragraphs
                        self.apply_styles(cell.paragraphs[0], paragraphs, custom_color)

                        self.document.save(self.file_path)
                        return True

        return False